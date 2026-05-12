import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from tempfile import TemporaryDirectory
from pathlib import Path

from docx import Document as DocxDocument

from app.main import app
from app.database import Base, get_db
from app import crud, schemas


# Use in-memory SQLite for tests
SQLALCHEMY_DATABASE_URL = "sqlite:///./test.db"

engine = create_engine(
    SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False}
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

Base.metadata.create_all(bind=engine)


def override_get_db():
    try:
        db = TestingSessionLocal()
        yield db
    finally:
        db.close()


app.dependency_overrides[get_db] = override_get_db

client = TestClient(app)


@pytest.fixture(autouse=True)
def setup_and_teardown():
    """Reset database before each test"""
    Base.metadata.create_all(bind=engine)
    yield
    Base.metadata.drop_all(bind=engine)


def test_create_user():
    """Test user creation"""
    db = TestingSessionLocal()
    user = crud.create_user(db, schemas.UserCreate(username="testuser", email="test@example.com"))
    assert user.username == "testuser"
    assert user.email == "test@example.com"
    db.close()


def test_create_document():
    """Test document creation"""
    db = TestingSessionLocal()
    user = crud.create_user(db, schemas.UserCreate(username="testuser", email="test@example.com"))
    doc = crud.create_document(
        db,
        schemas.DocumentCreate(title="Test Doc", content="Hello world"),
        user.id
    )
    assert doc.title == "Test Doc"
    assert doc.content == "Hello world"
    assert doc.owner_id == user.id
    db.close()


def test_share_document():
    """Test document sharing"""
    db = TestingSessionLocal()
    user1 = crud.create_user(db, schemas.UserCreate(username="alice", email="alice@example.com"))
    user2 = crud.create_user(db, schemas.UserCreate(username="bob", email="bob@example.com"))
    
    doc = crud.create_document(
        db,
        schemas.DocumentCreate(title="Shared Doc", content="Content"),
        user1.id
    )
    
    shared_doc = crud.share_document(db, doc.id, user2.id)
    assert user2 in shared_doc.shared_with_users
    db.close()


def test_api_create_document():
    """Test API endpoint for creating document"""
    # Create users first
    db = TestingSessionLocal()
    user = crud.create_user(db, schemas.UserCreate(username="alice", email="alice@example.com"))
    db.close()
    
    response = client.post(
        "/api/documents",
        json={"title": "API Test Doc", "content": "Test content"},
        headers={"x-user-id": str(user.id)}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["title"] == "API Test Doc"
    assert data["content"] == "Test content"


def test_api_list_documents():
    """Test API endpoint for listing documents"""
    db = TestingSessionLocal()
    user = crud.create_user(db, schemas.UserCreate(username="alice", email="alice@example.com"))
    user_id = user.id  # Store user_id before closing session
    doc = crud.create_document(
        db,
        schemas.DocumentCreate(title="Doc 1", content="Content 1"),
        user_id
    )
    db.close()
    
    response = client.get(
        "/api/documents",
        headers={"x-user-id": str(user_id)}
    )
    assert response.status_code == 200
    data = response.json()
    assert len(data) == 1
    assert data[0]["title"] == "Doc 1"
    assert data[0]["is_owner"] == True


def test_api_authentication_required():
    """Test that authentication is required"""
    response = client.get("/api/documents")
    assert response.status_code == 401


def test_api_upload_docx():
    """Test uploading a .docx file creates a document"""
    db = TestingSessionLocal()
    user = crud.create_user(db, schemas.UserCreate(username="alice", email="alice@example.com"))
    user_id = user.id
    db.close()

    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir) / "sample.docx"
        docx = DocxDocument()
        docx.add_heading("Imported Title", level=1)
        docx.add_paragraph("Imported body text.")
        docx.save(temp_path)

        with open(temp_path, "rb") as file_handle:
            response = client.post(
                "/api/documents/upload/file",
                headers={"x-user-id": str(user_id)},
                files={"file": ("sample.docx", file_handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

    assert response.status_code == 200
    data = response.json()
    assert data["title"] == "sample"
    assert "Imported Title" in data["content"]
    assert "Imported body text." in data["content"]
